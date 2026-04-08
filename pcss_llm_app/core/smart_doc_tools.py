"""
Smart Document Intelligence Tools — Docling-powered visual document parsing,
image extraction, and layout-preserving translation.

This module is OPTIONAL. It requires 'docling' and 'pillow' to be installed:
    pip install docling pillow

If these dependencies are not present, the tools will not be registered
with the agent, and the application will continue to function normally
using the standard read_pdf / read_docx tools.
"""

import os
import time
import base64
import mimetypes
import json
import ast
import re
from typing import List, Dict, Any, Optional
from openai import OpenAI

try:
    from langchain_core.tools import StructuredTool
except ImportError:
    from langchain.tools import StructuredTool

from pydantic import BaseModel, Field
from openai import OpenAI

# ---------------------------------------------------------------------------
# Docling imports — guarded to allow graceful fallback
# ---------------------------------------------------------------------------
try:
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.datamodel.base_models import InputFormat
    from docling_core.types.doc import PictureItem, TableItem

    DOCLING_AVAILABLE = True
except ImportError:
    DOCLING_AVAILABLE = False

# Re-use workspace sandbox from the main tools module
from pcss_llm_app.core.tools import _safe_path, _WorkspaceMixin


# ---------------------------------------------------------------------------
# ModelSelector — automatic best-model selection
# ---------------------------------------------------------------------------

class ModelSelector:
    """
    Selects the optimal LLM for a given task from the list of models
    currently available on the configured server.
    Not tied to any single model — queries /v1/models at runtime.
    """

    # Models known to support vision (image_url content type)
    VISION_PREFERENCE = [
        "Qwen3-VL-235B-A22B-Instruct",
    ]

    # Models ranked by translation quality (strongest first)
    TRANSLATION_PREFERENCE = [
        "Qwen3.5-397B-A17B",
        "DeepSeek-V3.1-vLLM",
        "DeepSeek-V3.1-vLLM-2",
        "GLM-4.7",
        "MiniMax-M2.5",
    ]

    def __init__(self, client: OpenAI):
        self.client = client
        self._available_models = None

    def _fetch_models(self) -> List[str]:
        """Fetch available models from the server (cached per instance)."""
        if self._available_models is None:
            try:
                models = self.client.models.list()
                self._available_models = [m.id for m in models.data]
            except Exception:
                self._available_models = []
        return self._available_models

    def select_vision_model(self) -> Optional[str]:
        """Return the best available vision model, or None."""
        available = self._fetch_models()
        for preferred in self.VISION_PREFERENCE:
            if preferred in available:
                return preferred
        # Fuzzy fallback — any model with "VL" or "vision" in name
        for m in available:
            ml = m.lower()
            if "vl" in ml or "vision" in ml or "multimodal" in ml:
                return m
        return None

    def select_translation_model(self) -> Optional[str]:
        """Return the strongest available text model for translation."""
        available = self._fetch_models()
        for preferred in self.TRANSLATION_PREFERENCE:
            if preferred in available:
                return preferred
        # Fallback: first available model
        return available[0] if available else None


# ---------------------------------------------------------------------------
# Pydantic schemas for tool arguments
# ---------------------------------------------------------------------------

class SmartReadDocumentSchema(BaseModel):
    file_path: str = Field(description="Path to the PDF or DOCX file (relative to workspace)")
    analyze_images: bool = Field(
        default=False,
        description="If True, each extracted image will be analyzed by a Vision LLM and its description appended to the output"
    )
    image_prompt: str = Field(
        default="Opisz szczegółowo zawartość tego obrazu. Jeśli to wykres lub tabela, podaj kluczowe dane.",
        description="Prompt sent to the Vision model for each image"
    )
    max_pages: Optional[int] = Field(
        default=None,
        description="Maximum number of pages to process (default: all)"
    )


class TranslateDocumentSchema(BaseModel):
    file_path: str = Field(description="Path to the PDF or DOCX file to translate (relative to workspace)")
    target_lang: str = Field(
        default="polski",
        description="Target language for translation (e.g., 'polski', 'English', 'Deutsch')"
    )
    source_lang: str = Field(
        default="auto",
        description="Source language (default: auto-detect)"
    )
    output_format: str = Field(
        default="docx",
        description="Output format: 'docx', 'html', or 'md'"
    )
    translation_model: Optional[str] = Field(
        default=None,
        description="Explicit model to use for translation. If None, auto-selects the strongest available model."
    )


class ExtractDocumentImagesSchema(BaseModel):
    file_path: str = Field(description="Path to the PDF or DOCX file (relative to workspace)")
    output_dir: str = Field(
        default="_extracted_images",
        description="Directory (relative to workspace) where extracted images will be saved"
    )


# ---------------------------------------------------------------------------
# SmartDocumentTools — the main tool class
# ---------------------------------------------------------------------------

class SmartDocumentTools(_WorkspaceMixin):
    """
    Docling-powered document intelligence tools that parse documents
    visually — understanding layout, tables, and images — just like
    a human reader would.

    Requires: pip install docling pillow
    """

    def __init__(self, root_dir: str, api_key: str, base_url: str = "https://llm.hpc.pcss.pl/v1", log_callback=None):
        self.root_dir = root_dir
        self.api_key = api_key
        self.base_url = base_url
        # Set a 120s timeout to handle slow model responses without hanging the agent indefinitely
        self.client = OpenAI(api_key=api_key, base_url=base_url, timeout=120.0)
        self.model_selector = ModelSelector(self.client)
        self.log_callback = log_callback

    def _log(self, message: str):
        """Helper to send logs to the main agent log if available."""
        if self.log_callback:
            self.log_callback(f"[SMART-DOC] {message}")

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _get_converter(self, generate_images: bool = False) -> "DocumentConverter":
        """Create a Docling DocumentConverter with appropriate options."""
        pipeline_options = PdfPipelineOptions()
        pipeline_options.generate_page_images = False  # We don't need full-page screenshots
        pipeline_options.generate_picture_images = generate_images
        pipeline_options.generate_table_images = False

        self._log("Initializing DocumentConverter... (First run will download AI models ~1.5GB)")
        return DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
            }
        )

    def _analyze_single_image(self, image_path: str, prompt: str, vision_model: str) -> str:
        """Send a single image to a Vision LLM for analysis."""
        try:
            mime_type, _ = mimetypes.guess_type(image_path)
            if not mime_type:
                mime_type = "image/png"

            with open(image_path, "rb") as f:
                b64 = base64.b64encode(f.read()).decode("utf-8")

            response = self.client.chat.completions.create(
                model=vision_model,
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {"type": "image_url", "image_url": {"url": f"data:{mime_type};base64,{b64}"}}
                    ]
                }]
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"[Vision analysis failed: {e}]"

    def _translate_text_block(self, text: str, source_lang: str, target_lang: str, model: str) -> str:
        """Translate a block of text using an LLM."""
        if not text.strip():
            return text

        lang_hint = f" from {source_lang}" if source_lang != "auto" else ""
        system_prompt = (
            f"You are a professional translator. Translate the following text{lang_hint} "
            f"to {target_lang}. Preserve ALL formatting: Markdown headings (#, ##), "
            f"bold (**), italic (*), lists, table syntax (|), and image references. "
            f"Do NOT add any commentary. Output ONLY the translated text."
        )

        self._log(f"→ Requesting translation ({len(text)} chars) from model '{model}'...")
        t_start = time.monotonic()
        try:
            response = self.client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": text}
                ],
                temperature=0.1
            )
            elapsed = time.monotonic() - t_start
            result = response.choices[0].message.content
            self._log(f"← Received translation ({len(result)} chars) in {elapsed:.2f}s")
            return result
        except Exception as e:
            elapsed = time.monotonic() - t_start
            self._log(f"✗ Translation failed after {elapsed:.2f}s: {e}")
            return f"[Translation error: {e}]\n{text}"

    # ------------------------------------------------------------------
    # Tool 1: smart_read_document
    # ------------------------------------------------------------------

    def smart_read_document(
        self,
        file_path: str,
        analyze_images: bool = False,
        image_prompt: str = "Opisz szczegółowo zawartość tego obrazu. Jeśli to wykres lub tabela, podaj kluczowe dane.",
        max_pages: Optional[int] = None,
    ) -> str:
        """
        Intelligently reads a PDF or DOCX file using AI-powered layout analysis (Docling).
        Unlike the basic read_pdf/read_docx tools, this tool:
        - Understands document structure: headings, paragraphs, tables, images
        - Extracts embedded images and saves them as separate PNG files
        - Optionally analyzes each image using a Vision LLM (analyze_images=True)
        - Preserves reading order even in complex multi-column layouts

        Args:
            file_path: Path to the PDF or DOCX file (relative to workspace).
            analyze_images: If True, each image will be analyzed by a Vision model.
            image_prompt: Prompt for image analysis (Polish/English).
            max_pages: Max pages to process (default: all).

        Returns:
            Rich Markdown with document content, table representations, and image descriptions.
        """
        try:
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File '{file_path}' not found in workspace."

            ext = os.path.splitext(file_path)[1].lower()
            if ext not in (".pdf", ".docx"):
                return f"Error: Unsupported format '{ext}'. Use .pdf or .docx."

            # Create image output directory
            images_dir_rel = os.path.join(
                os.path.dirname(file_path) or ".",
                f"_images_{os.path.splitext(os.path.basename(file_path))[0]}"
            )
            images_dir_full = self._get_full_path(images_dir_rel)
            os.makedirs(images_dir_full, exist_ok=True)

            # Run Docling converter
            converter = self._get_converter(generate_images=True)
            self._log(f"Starting layout analysis for '{file_path}'...")
            result = converter.convert(full_path)
            doc = result.document
            self._log(f"Document analysis complete. Found {len(list(doc.iterate_items()))} elements.")

            # Build output
            output_parts = [
                f"[SMART READ: {file_path}]",
                f"(Parsed with Docling AI layout analysis)\n"
            ]

            # Export to Markdown (Docling's native rich output)
            md_content = doc.export_to_markdown()

            # Extract and save images
            image_count = 0
            image_descriptions = []

            for idx, (item, _level) in enumerate(doc.iterate_items()):
                if isinstance(item, PictureItem):
                    image_count += 1
                    img_filename = f"figure_{image_count}.png"
                    img_path_rel = os.path.join(images_dir_rel, img_filename)
                    img_path_full = os.path.join(images_dir_full, img_filename)

                    # Save the image
                    try:
                        if hasattr(item, 'image') and item.image is not None:
                            if hasattr(item.image, 'pil_image') and item.image.pil_image is not None:
                                item.image.pil_image.save(img_path_full, "PNG")
                            else:
                                continue
                        else:
                            continue
                    except Exception:
                        continue

                    # Optionally analyze the image with Vision
                    if analyze_images:
                        vision_model = self.model_selector.select_vision_model()
                        if vision_model:
                            description = self._analyze_single_image(
                                img_path_full, image_prompt, vision_model
                            )
                            image_descriptions.append(
                                f"\n### 🖼 Figure {image_count} (`{img_path_rel}`)\n{description}\n"
                            )
                        else:
                            image_descriptions.append(
                                f"\n### 🖼 Figure {image_count} (`{img_path_rel}`)\n"
                                f"[No vision model available for analysis]\n"
                            )
                    else:
                        image_descriptions.append(
                            f"\n[IMAGE saved: {img_path_rel}]\n"
                        )

            # Truncate if needed
            if max_pages and len(md_content) > max_pages * 3000:
                md_content = md_content[:max_pages * 3000]
                md_content += "\n\n--- OUTPUT TRUNCATED (max_pages limit) ---"

            output_parts.append(md_content)

            if image_count > 0:
                output_parts.append(f"\n---\n## Extracted Images ({image_count} total)\n")
                output_parts.extend(image_descriptions)
                output_parts.append(
                    f"\nTIP: Use `analyze_image(file_path='{images_dir_rel}/figure_1.png')` "
                    f"to analyze any image individually."
                )

            return "\n".join(output_parts)

        except PermissionError:
            raise
        except Exception as e:
            return f"Error in smart_read_document: {str(e)}"

    # ------------------------------------------------------------------
    # Tool 2: translate_document
    # ------------------------------------------------------------------

    def translate_document(
        self,
        file_path: str,
        target_lang: str = "polski",
        source_lang: str = "auto",
        output_format: str = "docx",
        translation_model: Optional[str] = None,
    ) -> str:
        """
        Translates a PDF or DOCX document while preserving its structure and images.

        For DOCX files: performs in-place translation using python-docx, preserving
        original formatting (fonts, bold, italic, colors) and all images intact.

        For PDF files: uses Docling to extract structured Markdown + images,
        translates text blocks via LLM, then renders to the chosen output format
        (DOCX/HTML/Markdown) with images embedded.

        Args:
            file_path: Path to the PDF or DOCX file (relative to workspace).
            target_lang: Target language (e.g., 'polski', 'English', 'Deutsch').
            source_lang: Source language (default: 'auto' for auto-detection).
            output_format: Output format: 'docx', 'html', or 'md' (default: 'docx').
            translation_model: Explicit model name. If None, auto-selects best available.

        Returns:
            Path to the translated document file.
        """
        try:
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File '{file_path}' not found."

            ext = os.path.splitext(file_path)[1].lower()
            if ext not in (".pdf", ".docx"):
                return f"Error: Unsupported format '{ext}'. Use .pdf or .docx."

            # Select translation model
            model = translation_model or self.model_selector.select_translation_model()
            if not model:
                return "Error: No translation model available on the server."

            # Determine output path
            base_name = os.path.splitext(file_path)[0]
            out_ext = f".{output_format}" if not output_format.startswith(".") else output_format
            output_path = f"{base_name}_{target_lang}{out_ext}"

            if ext == ".docx":
                return self._translate_docx_inplace(
                    full_path, file_path, output_path, source_lang, target_lang, model
                )
            else:  # PDF
                self._log(f"Starting PDF processing via Docling: {file_path}")
                return self._translate_pdf_via_markdown(
                    full_path, file_path, output_path, source_lang, target_lang, model, output_format
                )

        except PermissionError:
            raise
        except Exception as e:
            return f"Error in translate_document: {str(e)}"

    def _translate_docx_inplace(
        self, full_path: str, rel_path: str, output_path: str,
        source_lang: str, target_lang: str, model: str
    ) -> str:
        """Translate DOCX in-place, preserving formatting and images."""
        from docx import Document

        doc = Document(full_path)
        translated_count = 0

        # Translate paragraphs
        total_paras = len(doc.paragraphs)
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue

            if i % 5 == 0:  # Log every 5 paragraphs
                self._log(f"Translating paragraph {i+1} of {total_paras}...")

            # Collect full paragraph text for context-aware translation
            original_text = para.text
            translated_text = self._translate_text_block(original_text, source_lang, target_lang, model)

            if translated_text and translated_text != original_text:
                # Strategy: if paragraph has a single run, replace directly.
                # If multiple runs, consolidate into first run to preserve its formatting.
                runs = para.runs
                if len(runs) == 1:
                    runs[0].text = translated_text
                elif len(runs) > 1:
                    # Keep first run's formatting, put all text there, clear the rest
                    runs[0].text = translated_text
                    for r in runs[1:]:
                        r.text = ""
                translated_count += 1

        # Translate table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue
                        original_text = para.text
                        translated_text = self._translate_text_block(
                            original_text, source_lang, target_lang, model
                        )
                        if translated_text and translated_text != original_text:
                            runs = para.runs
                            if len(runs) == 1:
                                runs[0].text = translated_text
                            elif len(runs) > 1:
                                runs[0].text = translated_text
                                for r in runs[1:]:
                                    r.text = ""
                            translated_count += 1

        # Save translated document
        output_full = self._get_full_path(output_path)
        os.makedirs(os.path.dirname(output_full) or ".", exist_ok=True)
        doc.save(output_full)

        self._log(f"✓ DOCX translation finished. Saved to: {output_path}")

        return (
            f"Successfully translated DOCX: {rel_path} → {output_path}\n"
            f"Translated {translated_count} text blocks to {target_lang}.\n"
            f"Model used: {model}\n"
            f"Original formatting and images preserved."
        )

    def _translate_pdf_via_markdown(
        self, full_path: str, rel_path: str, output_path: str,
        source_lang: str, target_lang: str, model: str, output_format: str
    ) -> str:
        """Translate PDF by extracting via Docling, translating Markdown, rebuilding."""

        # Step 1: Extract with Docling
        converter = self._get_converter(generate_images=True)
        result = converter.convert(full_path)
        doc = result.document

        # Step 2: Extract and save images
        images_dir_rel = os.path.join(
            os.path.dirname(rel_path) or ".",
            f"_images_{os.path.splitext(os.path.basename(rel_path))[0]}"
        )
        images_dir_full = self._get_full_path(images_dir_rel)
        os.makedirs(images_dir_full, exist_ok=True)

        image_count = 0
        for idx, (item, _level) in enumerate(doc.iterate_items()):
            if isinstance(item, PictureItem):
                image_count += 1
                img_filename = f"figure_{image_count}.png"
                img_path_full = os.path.join(images_dir_full, img_filename)
                try:
                    if hasattr(item, 'image') and item.image is not None:
                        if hasattr(item.image, 'pil_image') and item.image.pil_image is not None:
                            item.image.pil_image.save(img_path_full, "PNG")
                except Exception:
                    pass

        # Step 3: Get Markdown and translate in chunks
        md_content = doc.export_to_markdown()
        self._log("Layout extracted. Starting text translation in chunks...")

        # Split by double newline to translate logical blocks
        blocks = md_content.split("\n\n")
        total_blocks = len(blocks)
        translated_blocks = []
        batch = []
        batch_chars = 0
        MAX_BATCH_CHARS = 3000  # Translate in ~3k char chunks for quality

        for i, block in enumerate(blocks):
            if batch_chars + len(block) > MAX_BATCH_CHARS and batch:
                # Translate accumulated batch
                self._log(f"Translating block batch (up to block {i}/{total_blocks})...")
                batch_text = "\n\n".join(batch)
                translated = self._translate_text_block(batch_text, source_lang, target_lang, model)
                translated_blocks.append(translated)
                batch = []
                batch_chars = 0

            batch.append(block)
            batch_chars += len(block)

        # Translate remaining batch
        if batch:
            batch_text = "\n\n".join(batch)
            translated = self._translate_text_block(batch_text, source_lang, target_lang, model)
            translated_blocks.append(translated)

        translated_md = "\n\n".join(translated_blocks)

        # Step 4: Save in requested format
        if output_format == "md":
            output_full = self._get_full_path(output_path)
            os.makedirs(os.path.dirname(output_full) or ".", exist_ok=True)
            with open(output_full, "w", encoding="utf-8") as f:
                f.write(translated_md)
        elif output_format in ("docx", "html"):
            # Save as intermediate HTML then convert via Pandoc
            html_content = self._markdown_to_html(translated_md)
            if output_format == "html":
                output_full = self._get_full_path(output_path)
                os.makedirs(os.path.dirname(output_full) or ".", exist_ok=True)
                with open(output_full, "w", encoding="utf-8") as f:
                    f.write(html_content)
            else:
                # DOCX via pypandoc
                try:
                    import pypandoc
                    # Save HTML temp
                    html_temp = self._get_full_path(
                        os.path.splitext(output_path)[0] + "_temp.html"
                    )
                    with open(html_temp, "w", encoding="utf-8") as f:
                        f.write(html_content)

                    output_full = self._get_full_path(output_path)
                    os.makedirs(os.path.dirname(output_full) or ".", exist_ok=True)
                    pypandoc.convert_file(html_temp, "docx", outputfile=output_full)

                    # Clean up temp
                    try:
                        os.remove(html_temp)
                    except Exception:
                        pass
                except ImportError:
                    # Fallback: save as HTML instead
                    output_path = os.path.splitext(output_path)[0] + ".html"
                    output_full = self._get_full_path(output_path)
                    os.makedirs(os.path.dirname(output_full) or ".", exist_ok=True)
                    with open(output_full, "w", encoding="utf-8") as f:
                        f.write(html_content)

        self._log(f"✓ PDF translation finished. Saved to: {output_path}")

        return (
            f"Successfully translated PDF: {rel_path} → {output_path}\n"
            f"Extracted {image_count} images to {images_dir_rel}/\n"
            f"Model used: {model}\n"
            f"Target language: {target_lang}"
        )

    def _markdown_to_html(self, md_text: str) -> str:
        """Convert Markdown to a self-contained HTML document."""
        try:
            import markdown as md_lib
            body = md_lib.markdown(md_text, extensions=["tables", "fenced_code"])
        except ImportError:
            # Minimal fallback
            body = f"<pre>{md_text}</pre>"

        return f"""<!DOCTYPE html>
<html lang="pl">
<head>
    <meta charset="utf-8">
    <title>Translated Document</title>
    <style>
        body {{ font-family: 'Segoe UI', Tahoma, sans-serif; max-width: 900px; margin: 2em auto; padding: 0 1em; line-height: 1.6; }}
        table {{ border-collapse: collapse; width: 100%; margin: 1em 0; }}
        th, td {{ border: 1px solid #ccc; padding: 8px 12px; text-align: left; }}
        th {{ background: #f5f5f5; }}
        img {{ max-width: 100%; height: auto; }}
        h1, h2, h3 {{ color: #333; }}
    </style>
</head>
<body>
{body}
</body>
</html>"""

    # ------------------------------------------------------------------
    # Tool 3: extract_document_images
    # ------------------------------------------------------------------

    def extract_document_images(
        self,
        file_path: str,
        output_dir: str = "_extracted_images",
    ) -> str:
        """
        Extracts ALL images from a PDF or DOCX file and saves them as separate PNG files.
        Use this when you need to examine, analyze, or reuse individual images
        from a document.

        After extraction, you can use the `analyze_image` tool on any of the
        saved files to get a detailed AI description.

        Args:
            file_path: Path to the PDF or DOCX file (relative to workspace).
            output_dir: Directory where images will be saved (default: '_extracted_images').

        Returns:
            List of saved image file paths with their dimensions.
        """
        try:
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File '{file_path}' not found."

            ext = os.path.splitext(file_path)[1].lower()
            if ext not in (".pdf", ".docx"):
                return f"Error: Unsupported format '{ext}'. Use .pdf or .docx."

            # Create output directory
            output_dir_full = self._get_full_path(output_dir)
            os.makedirs(output_dir_full, exist_ok=True)

            # Parse with Docling (images enabled)
            converter = self._get_converter(generate_images=True)
            result = converter.convert(full_path)
            doc = result.document

            saved_images = []
            for idx, (item, _level) in enumerate(doc.iterate_items()):
                if isinstance(item, PictureItem):
                    try:
                        if hasattr(item, 'image') and item.image is not None:
                            pil_img = getattr(item.image, 'pil_image', None)
                            if pil_img is not None:
                                img_idx = len(saved_images) + 1
                                img_filename = f"image_{img_idx}.png"
                                img_path_full = os.path.join(output_dir_full, img_filename)
                                img_path_rel = os.path.join(output_dir, img_filename)

                                pil_img.save(img_path_full, "PNG")
                                w, h = pil_img.size
                                saved_images.append(f"  {img_path_rel}  ({w}×{h} px)")
                    except Exception as e:
                        saved_images.append(f"  [Failed to save image {idx}: {e}]")

            if not saved_images:
                return f"No images found in '{file_path}'."

            result_lines = [
                f"Extracted {len(saved_images)} image(s) from '{file_path}' to '{output_dir}/':",
                "",
            ] + saved_images + [
                "",
                f"TIP: Use analyze_image(file_path='{output_dir}/image_1.png') to analyze any image."
            ]

            return "\n".join(result_lines)

        except PermissionError:
            raise
        except Exception as e:
            return f"Error extracting images: {str(e)}"

    # ------------------------------------------------------------------
    # Tool registration
    # ------------------------------------------------------------------

    def get_tools(self) -> list:
        """Return list of LangChain StructuredTool instances."""
        return [
            StructuredTool.from_function(
                func=self.smart_read_document,
                name="smart_read_document",
                description=(
                    "Inteligentny odczyt dokumentu PDF/DOCX z analizą layoutu AI (Docling). "
                    "W porównaniu do read_pdf/read_docx: rozumie strukturę wielokolumnową, "
                    "wyciąga grafiki i zapisuje je jako osobne pliki PNG, opcjonalnie analizuje "
                    "grafiki modelem Vision. Użyj gdy potrzebujesz pełnej analizy dokumentu "
                    "z grafikami. Args: file_path, analyze_images=False, image_prompt, max_pages."
                ),
                args_schema=SmartReadDocumentSchema,
            ),
            StructuredTool.from_function(
                func=self.translate_document,
                name="translate_document",
                description=(
                    "Tłumaczy dokument PDF lub DOCX z zachowaniem struktury, formatowania i grafik. "
                    "Dla DOCX: tłumaczenie in-place z zachowaniem czcionek, boldów i kolorów. "
                    "Dla PDF: ekstrakcja Docling → tłumaczenie LLM → eksport do DOCX/HTML/MD. "
                    "Automatycznie dobiera najsilniejszy dostępny model do tłumaczenia. "
                    "Args: file_path, target_lang='polski', source_lang='auto', "
                    "output_format='docx', translation_model=None."
                ),
                args_schema=TranslateDocumentSchema,
            ),
            StructuredTool.from_function(
                func=self.extract_document_images,
                name="extract_document_images",
                description=(
                    "Wyciąga WSZYSTKIE grafiki z pliku PDF lub DOCX i zapisuje je jako osobne "
                    "pliki PNG w workspace. Po ekstrakcji możesz użyć analyze_image na dowolnym "
                    "z nich. Args: file_path, output_dir='_extracted_images'."
                ),
                args_schema=ExtractDocumentImagesSchema,
            ),
        ]
