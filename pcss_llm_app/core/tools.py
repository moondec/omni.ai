import os
import sys
import io
import contextlib
import subprocess
import shlex
import time
import re
import json
import datetime as dt
from typing import Optional, List
try:
    from langchain_core.tools import tool, StructuredTool
except ImportError:
    from langchain.tools import tool, StructuredTool
from pydantic import BaseModel, Field
from docx import Document
from pypdf import PdfReader
from openai import OpenAI
import base64
import mimetypes
try:
    import pypandoc
except ImportError:
    pypandoc = None

import re
import xml.dom.minidom

try:
    from ddgs import DDGS
except ImportError:
    from duckduckgo_search import DDGS


# ---------------------------------------------------------------------------
# Workspace sandbox — all tools must call _safe_path() before touching disk.
# ---------------------------------------------------------------------------

def _safe_path(root_dir: str, user_path: str) -> str:
    """
    Resolve a user-supplied path relative to root_dir and verify it stays
    within the workspace.  Raises PermissionError on sandbox violations.

    Protects against:
    - Path traversal:  ../../etc/passwd
    - Absolute paths:  /etc/passwd  (explicitly rejected)
    - Symlink escapes: symlinks that resolve outside the workspace
    """
    root = os.path.realpath(root_dir)

    # Explicitly reject absolute paths — do not silently reroot them
    if os.path.isabs(user_path):
        raise PermissionError(
            f"Access denied: absolute path '{user_path}' is not allowed. "
            f"Provide a path relative to the workspace."
        )

    candidate = os.path.realpath(os.path.join(root, user_path))

    # os.path.commonpath raises ValueError if paths are on different drives (Windows)
    try:
        common = os.path.commonpath([root, candidate])
    except ValueError:
        common = ""

    if common != root:
        raise PermissionError(
            f"Access denied: '{user_path}' resolves outside the workspace "
            f"(workspace: '{root}', resolved: '{candidate}')."
        )
    return candidate


class _WorkspaceMixin:
    """Mixin that provides the safe path resolver for tool classes."""
    root_dir: str

    def _get_full_path(self, file_path: str) -> str:
        return _safe_path(self.root_dir, file_path)


# ---------------------------------------------------------------------------
# Pydantic schemas for tools with multiple parameters
# ---------------------------------------------------------------------------
class SaveDocumentSchema(BaseModel):
    file_path: str = Field(description="Target file name with extension (e.g., 'report.pdf', 'summary.docx')")
    content: str = Field(description="HTML-formatted content (use <h1>, <p>, <ul>, <li>, <b>, <i> tags for formatting)")
    title: str = Field(default="Document", description="Document title (used in HTML header)")

class DocumentTools(_WorkspaceMixin):
    def __init__(self, root_dir: str):
        self.root_dir = root_dir

    def _get_full_path(self, file_path: str) -> str:
        try:
            return _safe_path(self.root_dir, file_path)
        except PermissionError as e:
            raise PermissionError(str(e))
    
    # @tool("write_file")
    def write_file(self, file_path: str, text: str):
        """
        Creates or overwrites a file with the given content. 
        Features:
        - Auto-formats XML files (.xml) with indentation.
        - Automatically fixes literal '\\n' characters.
        """
        try:
            full_path = self._get_full_path(file_path)
            
            # Ensure parent directories exist
            os.makedirs(os.path.dirname(full_path), exist_ok=True)

            # Auto-correction for literal newlines if passed as string literal
            if "\\n" in text and "\n" not in text:
                 text = text.replace("\\n", "\n")

            # Logic for XML Formatting
            if file_path.lower().endswith(".xml"):
                try:
                    # Parse and pretty print
                    # Remove whitespace between tags to prevents weird spacing before pretty-printing
                    clean_text = "".join(line.strip() for line in text.split('\n'))
                    xml_parsed = xml.dom.minidom.parseString(clean_text)
                    final_content = xml_parsed.toprettyxml(indent="  ")
                    # Fix extra newlines sometimes introduced by minidom text nodes
                    final_content = "\n".join([line for line in final_content.split('\n') if line.strip()])
                except Exception as xml_err:
                    # Fallback if XML is malformed
                    # print(f"Warning: XML formatting failed ({xml_err}). Saving as raw text.")
                    final_content = text
            else:
                final_content = text

            with open(full_path, "w", encoding="utf-8") as f:
                f.write(final_content)
                
            return f"Successfully saved file: {file_path}"
        except PermissionError:
            raise
        except Exception as e:
            return f"Error writing file: {str(e)}"

    # @tool("write_docx")
    def write_docx(self, file_path: str, text: str):
        """
        Creates a new Word document (.docx) with the given text.
        Args:
            file_path: The name of the file to save (e.g., 'document.docx').
            text: The text content to write into the document.
        """
        try:
            full_path = self._get_full_path(file_path)
            doc = Document()
            for line in text.split('\n'):
                doc.add_paragraph(line)
            doc.save(full_path)
            return f"Successfully created DOCX file: {file_path}"
        except PermissionError:
            raise
        except Exception as e:
            return f"Error writing DOCX file: {str(e)}"

    # @tool("read_docx")
    def read_docx(self, file_path: str, para_start: int = 1, para_end: int = None) -> str:
        """
        Reads a Word document (.docx) preserving structure: headings, tables, images.
        Args:
            file_path: The name of the file to read.
            para_start: First block to read (1-indexed, counts paragraphs+tables+images).
            para_end: Last block to read inclusive (default: last block).
                      Read in chunks of 100-150 blocks for large documents.
        Output format:
            - Headings  → ## Heading text  (# H1, ## H2, etc.)
            - Tables    → Markdown pipe tables
            - Images    → [IMAGE: inline image, ~WxH px] placeholder
            - Paragraphs → plain text
        """
        try:
            from docx.oxml.ns import qn
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File {file_path} not found."

            doc = Document(full_path)
            body = doc.element.body

            # ── Walk body in document order ──────────────────────────────
            # Each child is a paragraph (w:p), table (w:tbl), or section (w:sectPr).
            blocks = []  # list of rendered strings

            for child in body:
                local = child.tag.split('}')[-1] if '}' in child.tag else child.tag

                # ── Paragraph ────────────────────────────────────────────
                if local == 'p':
                    from docx.text.paragraph import Paragraph as DocxParagraph
                    para = DocxParagraph(child, doc)
                    style_name = para.style.name if para.style else ""

                    # Detect heading level
                    heading_level = 0
                    if style_name.startswith("Heading"):
                        try:
                            heading_level = int(style_name.split()[-1])
                        except ValueError:
                            heading_level = 1

                    text = para.text.strip()

                    # Detect inline images in this paragraph
                    has_image = child.find('.//' + qn('a:blip')) is not None or \
                                child.find('.//' + qn('v:imagedata')) is not None
                    if has_image:
                        # Try to get EMU dimensions
                        ext_elem = child.find('.//' + qn('wp:extent'))
                        if ext_elem is not None:
                            cx = int(ext_elem.get('cx', 0))
                            cy = int(ext_elem.get('cy', 0))
                            # 1 EMU = 1/914400 inch, 96dpi
                            w_px = round(cx / 914400 * 96)
                            h_px = round(cy / 914400 * 96)
                            img_info = f"[IMAGE: inline image, ~{w_px}×{h_px} px]"
                        else:
                            img_info = "[IMAGE: inline image]"
                        blocks.append(img_info + (f"\n{text}" if text else ""))
                        continue

                    if not text:
                        continue  # skip blank paragraphs

                    if heading_level:
                        prefix = "#" * min(heading_level, 6)
                        blocks.append(f"{prefix} {text}")
                    else:
                        blocks.append(text)

                # ── Table ────────────────────────────────────────────────
                elif local == 'tbl':
                    from docx.table import Table as DocxTable
                    tbl = DocxTable(child, doc)
                    rows = tbl.rows
                    if not rows:
                        continue
                    md_rows = []
                    for r_idx, row in enumerate(rows):
                        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                        md_rows.append("| " + " | ".join(cells) + " |")
                        if r_idx == 0:
                            md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
                    blocks.append("\n".join(md_rows))

                # ── Section properties (skip) ─────────────────────────────
                # else: ignore w:sectPr etc.

            total_blocks = len(blocks)

            # ── Paginate ─────────────────────────────────────────────────
            start_idx = max(0, (para_start or 1) - 1)
            end_idx   = min(total_blocks, para_end or total_blocks)

            if start_idx >= total_blocks:
                return f"Error: para_start={para_start} exceeds total blocks ({total_blocks})."

            chunk = blocks[start_idx:end_idx]
            header = (
                f"[DOCX: {file_path} | Blocks {start_idx+1}–{end_idx} of {total_blocks}]\n"
                f"(Blocks = paragraphs + tables + images, in document order)\n"
            )
            return header + "\n\n".join(chunk)

        except PermissionError:
            raise
        except Exception as e:
            return f"Error reading DOCX file: {str(e)}"


    # @tool("read_pdf")
    def read_pdf(self, file_path: str, page_start: int = 1, page_end: int = None) -> str:
        """
        Reads a PDF file with structure: text, tables (Markdown), and image placeholders.
        Args:
            file_path: The name of the file to read.
            page_start: First page to read (1-indexed, default: 1).
            page_end: Last page to read inclusive (1-indexed, default: last page).
                      For large PDFs, read in chunks of 5-10 pages to avoid context overflow.
        Output format per page:
            - Text paragraphs as plain text
            - Tables as Markdown pipe tables
            - Images as [IMAGE: ~WxH px] placeholders
        """
        try:
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File {file_path} not found."

            reader = PdfReader(full_path)
            total_pages = len(reader.pages)

            start_idx = max(0, (page_start or 1) - 1)
            end_idx   = min(total_pages, (page_end or total_pages))

            if start_idx >= total_pages:
                return f"Error: page_start={page_start} exceeds total pages ({total_pages})."

            # ── Try pdfplumber for rich extraction ───────────────────────
            try:
                import pdfplumber

                output_parts = [
                    f"[PDF: {file_path} | Pages {start_idx+1}–{end_idx} of {total_pages}]"
                ]

                with pdfplumber.open(full_path) as pdf:
                    for page_num in range(start_idx, end_idx):
                        page = pdf.pages[page_num]
                        output_parts.append(f"\n── Page {page_num + 1} ──")

                        # ── Images: use pdfplumber .images (has x0/y0/x1/y1 bbox) ──
                        images_on_page = []
                        try:
                            for img in page.images:  # pdfplumber image objects
                                w = round(abs(img.get('x1', 0) - img.get('x0', 0)))
                                h = round(abs(img.get('y1', 0) - img.get('y0', 0)))
                                name = img.get('name', 'image')
                                images_on_page.append(f"[IMAGE: '{name}', ~{w}×{h} pt]")
                        except Exception:
                            pass
                        for img_ph in images_on_page:
                            output_parts.append(img_ph)

                        # ── Tables: extract and render as Markdown ──────────
                        tables = page.extract_tables()
                        table_regions = []
                        for tbl in tables:
                            if not tbl:
                                continue
                            md_rows = []
                            for r_idx, row in enumerate(tbl):
                                cells = [str(c).strip() if c else "" for c in row]
                                md_rows.append("| " + " | ".join(cells) + " |")
                                if r_idx == 0:
                                    md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
                            table_md = "\n".join(md_rows)
                            table_regions.append(table_md)

                        # ── Text: extract without table bounding boxes ──────
                        # Remove table regions so text outside tables is clean
                        if tables:
                            page_no_tables = page
                            try:
                                table_bboxes = [tbl.bbox for tbl in page.find_tables()]
                                page_no_tables = page.filter(
                                    lambda obj: not any(
                                        obj.get("x0", 0) >= bbox[0] - 2
                                        and obj.get("x1", 0) <= bbox[2] + 2
                                        and obj.get("top", 0) >= bbox[1] - 2
                                        and obj.get("bottom", 0) <= bbox[3] + 2
                                        for bbox in table_bboxes
                                    )
                                )
                                page_text = page_no_tables.extract_text() or ""
                            except Exception:
                                page_text = page.extract_text() or ""
                        else:
                            page_text = page.extract_text() or ""

                        if page_text.strip():
                            output_parts.append(page_text.strip())

                        for tbl_md in table_regions:
                            output_parts.append("\n" + tbl_md)

                return "\n".join(output_parts)

            # ── Fallback: plain pypdf extraction ─────────────────────────
            except ImportError:
                output_parts = [
                    f"[PDF: {file_path} | Pages {start_idx+1}–{end_idx} of {total_pages}]"
                ]
                for page_num in range(start_idx, end_idx):
                    pypdf_page = reader.pages[page_num]
                    output_parts.append(f"\n── Page {page_num + 1} ──")
                    # Images via pypdf (name only, no dimensions)
                    try:
                        for img in pypdf_page.images:
                            output_parts.append(f"[IMAGE: '{img.name}']")
                    except Exception:
                        pass
                    text = pypdf_page.extract_text() or ""
                    if text.strip():
                        output_parts.append(text.strip())
                return "\n".join(output_parts)

        except PermissionError:
            raise
        except Exception as e:
            return f"Error reading PDF file: {str(e)}"


    # @tool("read_xlsx")
    def read_xlsx(
        self,
        file_path: str,
        sheet: str = None,
        row_start: int = 1,
        row_end: int = None,
    ) -> str:
        """
        Reads an Excel file (.xlsx) and renders a sheet as a Markdown table.
        Args:
            file_path: The name of the .xlsx file to read.
            sheet: Sheet name to read (default: first sheet). Use list_sheets=True
                   call (sheet='?') to get all sheet names first.
            row_start: First data row to read, counting from 1 incl. header (default: 1).
            row_end: Last row to read inclusive (default: row_start + 49, i.e. 50 rows).
                     For large sheets read in chunks of 50-100 rows.
        Output:
            - Lists available sheet names in the header
            - Renders rows as a Markdown pipe table
            - Row 1 is used as the column header
        """
        try:
            import openpyxl
            from openpyxl.utils import get_column_letter

            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File {file_path} not found."

            wb = openpyxl.load_workbook(full_path, read_only=True, data_only=True)
            sheet_names = wb.sheetnames

            # Special call: sheet='?' → just list sheets
            if sheet == '?':
                return (
                    f"[XLSX: {file_path}]\n"
                    f"Available sheets ({len(sheet_names)}): "
                    + ", ".join(f"'{s}'" for s in sheet_names)
                )

            # Resolve sheet
            if sheet is None:
                ws = wb.active
                sheet_name = ws.title
            elif sheet in sheet_names:
                ws = wb[sheet]
                sheet_name = sheet
            else:
                return (
                    f"Error: Sheet '{sheet}' not found. "
                    f"Available: {sheet_names}"
                )

            # Read all rows into memory (read_only streams, so collect first)
            all_rows = list(ws.iter_rows(values_only=True))
            total_rows = len(all_rows)

            if total_rows == 0:
                return f"[XLSX: {file_path} | Sheet: '{sheet_name}'] Sheet is empty."

            # Resolve row range (1-indexed)
            r_start = max(1, row_start or 1)
            r_end   = min(total_rows, row_end or (r_start + 49))  # default 50 rows

            chunk = all_rows[r_start - 1 : r_end]  # 0-indexed slice

            # Build Markdown table
            def cell_str(v):
                if v is None:
                    return ""
                import datetime
                if isinstance(v, datetime.datetime):
                    return v.strftime("%Y-%m-%d %H:%M")
                if isinstance(v, datetime.date):
                    return v.strftime("%Y-%m-%d")
                return str(v).strip().replace("|", "\\|").replace("\n", " ")

            rows_md = []
            for r_idx, row in enumerate(chunk):
                cells = [cell_str(c) for c in row]
                rows_md.append("| " + " | ".join(cells) + " |")
                # Header separator after first row
                if r_idx == 0:
                    rows_md.append("| " + " | ".join(["---"] * len(cells)) + " |")

            remaining = total_rows - r_end
            footer = ""
            if remaining > 0:
                next_end = min(total_rows, r_end + (r_end - r_start + 1))
                footer = (
                    f"\n[{remaining} more rows. Call with "
                    f"row_start={r_end + 1}, row_end={next_end} to continue.]"
                )

            header = (
                f"[XLSX: {file_path} | Sheet: '{sheet_name}' "
                f"({len(sheet_names)} sheets total) | "
                f"Rows {r_start}–{r_end} of {total_rows}]\n"
                f"All sheets: {', '.join(sheet_names)}\n"
            )
            wb.close()
            return header + "\n".join(rows_md) + footer

        except PermissionError:
            raise
        except Exception as e:
            return f"Error reading XLSX file: {str(e)}"

    def save_document(self, file_path: str, content: str, title: str = "Document") -> str:
        """
        Saves formatted content to a document file. Supports: .pdf, .docx, .html, .txt
        For PDF/DOCX: automatically creates HTML first, then converts via Pandoc.
        Remote images are automatically downloaded and embedded.
        Args:
            file_path: Target file name with extension (e.g., 'report.pdf', 'summary.docx').
            content: HTML-formatted content (use <h1>, <p>, <ul>, <li>, <b>, <i> tags for formatting).
            title: Document title (used in HTML header).
        """
        import subprocess
        import tempfile
        import re
        import requests
        import hashlib
        
        ext = os.path.splitext(file_path)[1].lower()
        full_path = self._get_full_path(file_path)
        
        # Helper function to download remote images and replace with local paths
        def download_and_embed_images(html: str, temp_dir: str) -> str:
            """Download remote images and replace URLs with local file paths."""
            img_pattern = r'<img\s+[^>]*src=["\']([^"\']+)["\'][^>]*>'
            
            def replace_img(match):
                img_tag = match.group(0)
                src = match.group(1)
                
                # Skip local/relative paths and data URIs
                if not src.startswith(('http://', 'https://')) or src.startswith('data:'):
                    return img_tag
                
                try:
                    # Download image
                    response = requests.get(src, timeout=15, headers={
                        'User-Agent': 'Mozilla/5.0 (compatible; DocumentGenerator/1.0)'
                    })
                    response.raise_for_status()
                    
                    # Determine file extension from content-type or URL
                    content_type = response.headers.get('content-type', '')
                    if 'png' in content_type:
                        img_ext = '.png'
                    elif 'gif' in content_type:
                        img_ext = '.gif'
                    elif 'webp' in content_type:
                        img_ext = '.webp'
                    else:
                        img_ext = '.jpg'
                    
                    # Create unique filename
                    img_hash = hashlib.md5(src.encode()).hexdigest()[:12]
                    local_filename = f"img_{img_hash}{img_ext}"
                    local_path = os.path.join(temp_dir, local_filename)
                    
                    # Save image
                    with open(local_path, 'wb') as img_file:
                        img_file.write(response.content)
                    
                    # Replace src in tag
                    new_tag = img_tag.replace(src, local_path)
                    return new_tag
                    
                except Exception as e:
                    # If download fails, keep original and add alt text
                    # print(f"Warning: Could not download image {src}: {e}")
                    return img_tag
            
            return re.sub(img_pattern, replace_img, html, flags=re.IGNORECASE)
        
        # Wrap content in proper HTML structure if not already
        if not content.strip().startswith('<!DOCTYPE') and not content.strip().startswith('<html'):
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        h1 {{ color: #333; }}
        h2 {{ color: #555; }}
        ul, ol {{ margin-left: 20px; }}
        li {{ margin-bottom: 10px; }}
        .source {{ font-style: italic; color: #666; }}
        img {{ max-width: 100%; height: auto; }}
    </style>
</head>
<body>
{content}
</body>
</html>"""
        else:
            html_content = content
        
        try:
            # For TXT: just write plain text (strip HTML)
            if ext == '.txt':
                from html import unescape
                import re
                plain_text = re.sub(r'<[^>]+>', '', html_content)
                plain_text = unescape(plain_text)
                with open(full_path, 'w', encoding='utf-8') as f:
                    f.write(plain_text)
                return f"Successfully saved text file: {file_path}"
            
            # For HTML: write directly
            if ext == '.html':
                with open(full_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                return f"Successfully saved HTML file: {file_path}"
            
            # For PDF and DOCX: create temp HTML, then convert
            if ext in ['.pdf', '.docx']:
                # Create temp directory for images
                temp_img_dir = tempfile.mkdtemp(prefix='docgen_images_')
                
                # Download remote images and get modified HTML
                html_with_local_images = download_and_embed_images(html_content, temp_img_dir)
                
                # Create temp HTML file
                html_path = full_path.replace(ext, '.html')
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(html_with_local_images)
                
                # Try conversion
                if ext == '.pdf':
                    # Method 1: weasyprint (best for HTML->PDF, pure Python)
                    try:
                        from weasyprint import HTML
                        HTML(filename=html_path).write_pdf(full_path)
                        if os.path.exists(full_path) and os.path.getsize(full_path) > 0:
                            return f"Successfully saved PDF: {file_path}"
                    except ImportError:
                        pass  # weasyprint not installed
                    except Exception as e:
                        # print(f"weasyprint failed: {e}")
                        pass
                    
                    # Method 2: wkhtmltopdf (fallback)
                    try:
                        result = subprocess.run(
                            ['wkhtmltopdf', '--encoding', 'utf-8', '--quiet', html_path, full_path],
                            capture_output=True, text=True, timeout=60
                        )
                        if result.returncode == 0 and os.path.exists(full_path):
                            return f"Successfully saved PDF: {file_path}"
                    except FileNotFoundError:
                        pass  # wkhtmltopdf not installed
                    except Exception:
                        pass
                    
                    # Method 3: Pandoc with wkhtmltopdf engine
                    if pypandoc:
                        try:
                            pypandoc.convert_file(html_path, 'pdf', outputfile=full_path, 
                                                  extra_args=['--pdf-engine=wkhtmltopdf'])
                            if os.path.exists(full_path) and os.path.getsize(full_path) > 0:
                                return f"Successfully saved PDF: {file_path}"
                        except Exception:
                            pass
                    
                    # Fallback: Create DOCX instead and inform user
                    docx_path = full_path.replace('.pdf', '.docx')
                    if pypandoc:
                        try:
                            pypandoc.convert_file(html_path, 'docx', outputfile=docx_path)
                            return f"PDF conversion failed (missing wkhtmltopdf). Created DOCX instead: {file_path.replace('.pdf', '.docx')}. Install wkhtmltopdf: 'brew install wkhtmltopdf'"
                        except Exception as e:
                            return f"Error: Could not create PDF or DOCX. Details: {str(e)}"
                    return "Error: No PDF conversion tool available. Install wkhtmltopdf or pypandoc."
                
                elif ext == '.docx':
                    if pypandoc:
                        try:
                            pypandoc.convert_file(html_path, 'docx', outputfile=full_path)
                            if os.path.exists(full_path):
                                return f"Successfully saved DOCX: {file_path}"
                        except Exception as e:
                            return f"Error converting to DOCX: {str(e)}"
                    else:
                        # Fallback: use python-docx directly
                        try:
                            from bs4 import BeautifulSoup
                            from docx import Document as DocxDocument
                            
                            soup = BeautifulSoup(html_content, 'html.parser')
                            doc = DocxDocument()
                            doc.add_heading(title, 0)
                            
                            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'li']):
                                text = element.get_text(strip=True)
                                if element.name.startswith('h'):
                                    level = int(element.name[1])
                                    doc.add_heading(text, level)
                                else:
                                    doc.add_paragraph(text)
                            
                            doc.save(full_path)
                            return f"Successfully saved DOCX: {file_path}"
                        except Exception as e:
                            return f"Error creating DOCX: {str(e)}"
            
            return f"Unsupported format: {ext}. Use .pdf, .docx, .html, or .txt"
            
        except PermissionError:
            raise
        except Exception as e:
            return f"Error saving document: {str(e)}"

    def get_tools(self):
        """Returns the list of bound tools."""
        # We need to bind self to the tool methods manually or re-create them?
        # The @tool decorator works on functions. If used on methods, self needs handling.
        # Better approach: Define these as unbound functions or use StructuredTool.from_function
        # wrapping the instance methods.
        
        # Let's use closure-based definition inside get_tools to capture self
        # Or better, just use the methods since I used @tool decorator?
        # Actually @tool on methods might be tricky with 'self' in signature for LangChain.
        
        # Simpler Pattern:
        from langchain_core.tools import StructuredTool

        return [
            StructuredTool.from_function(
                func=self.save_document,
                name="save_document",
                description="Saves formatted content to a document file (.pdf, .docx, .html, .txt). Use HTML tags for formatting (h1, h2, p, ul, li, b, i). For PDF/DOCX: automatically handles conversion. This is the PREFERRED tool for creating formatted documents.",
                args_schema=SaveDocumentSchema
            ),
            StructuredTool.from_function(
                func=self.write_docx,
                name="write_docx",
                description="Creates a new Word document (.docx) with plain text. Use save_document instead for formatted content."
            ),
            StructuredTool.from_function(
                func=self.write_file,
                name="write_file",
                description="Creates or overwrites a file. PREFERRED for source code, XML, JSON, or plain text. Auto-formats XML."
            ),
            StructuredTool.from_function(
                func=self.read_docx,
                name="read_docx",
                description=(
                    "Reads text from a Word document (.docx) paragraph by paragraph. "
                    "Args: file_path (str), para_start (int, default 1), para_end (int, default last). "
                    "IMPORTANT: Large DOCX files MUST be read in chunks of 100-200 paragraphs. "
                    "Example: {\"file_path\": \"doc.docx\", \"para_start\": 1, \"para_end\": 150}. "
                    "The response includes total paragraph count so you know how many chunks remain."
                )
            ),
            StructuredTool.from_function(
                func=self.read_pdf,
                name="read_pdf",
                description=(
                    "Reads text from a PDF file page by page. "
                    "Args: file_path (str), page_start (int, default 1), page_end (int, default last page). "
                    "IMPORTANT: Large PDFs MUST be read in chunks of 5-10 pages to avoid context overflow. "
                    "Example: {\"file_path\": \"paper.pdf\", \"page_start\": 1, \"page_end\": 10}. "
                    "The response includes total page count so you know how many chunks remain."
                )
            ),
            StructuredTool.from_function(
                func=self.read_xlsx,
                name="read_xlsx",
                description=(
                    "Reads an Excel spreadsheet (.xlsx) and renders it as a Markdown table. "
                    "Args: file_path (str), sheet (str, default: first sheet — pass '?' to list all sheets), "
                    "row_start (int, default 1), row_end (int, default row_start+49). "
                    "IMPORTANT: Read in chunks of 50-100 rows for large sheets. "
                    "First call sheet='?' to discover sheet names. "
                    "Example: {\"file_path\": \"data.xlsx\", \"sheet\": \"Results\", \"row_start\": 1, \"row_end\": 50}."
                )
            ),
        ]


class CreateDirectorySchema(BaseModel):
    dir_path: str = Field(description="The name or path of the directory to create.")

class ListDirectorySchema(BaseModel):
    dir_path: str = Field(default=".", description="The path of the directory to list.")

class FolderTools(_WorkspaceMixin):
    def __init__(self, root_dir: str):
        self.root_dir = root_dir

    def create_directory(self, dir_path: str) -> str:
        """
        Creates a new directory (folder) at the specified path.
        Args:
            dir_path: The name or path of the directory to create.
        """
        try:
            full_path = self._get_full_path(dir_path)
            os.makedirs(full_path, exist_ok=True)
            return f"Successfully created directory: {dir_path}"
        except PermissionError:
            raise
        except Exception as e:
            return f"Error creating directory: {str(e)}"

    def list_directory(self, dir_path: str = ".") -> str:
        """
        Lists files and subdirectories in the specified path.
        Args:
            dir_path: Path to list (default is root).
        """
        try:
            full_path = self._get_full_path(dir_path)
            if not os.path.exists(full_path):
                return f"Error: Directory '{dir_path}' does not exist."
            
            items = os.listdir(full_path)
            output = []
            
            # Sort items: directories first, then files
            dims = []
            files = []
            
            for item in items:
                if item.startswith('.'): continue # Skip hidden files
                item_path = os.path.join(full_path, item)
                if os.path.isdir(item_path):
                    dims.append(item)
                else:
                    files.append(item)
            
            dims.sort()
            files.sort()
            
            if not dims and not files:
                return f"Directory '{dir_path}' is empty."

            for d in dims:
                 output.append(f"[DIR]  {d}")
            
            for f in files:
                try:
                    size_bytes = os.path.getsize(os.path.join(full_path, f))
                    # simple size formatting
                    if size_bytes < 1024:
                        size_str = f"{size_bytes} B"
                    elif size_bytes < 1024 * 1024:
                        size_str = f"{size_bytes/1024:.1f} KB"
                    else:
                        size_str = f"{size_bytes/(1024*1024):.1f} MB"
                    output.append(f"[FILE] {f} ({size_str})")
                except:
                    output.append(f"[FILE] {f}")
            
            return "\n".join(output)
        except PermissionError:
            raise
        except Exception as e:
            return f"Error listing directory: {str(e)}"

    def get_tools(self):
        from langchain_core.tools import StructuredTool

        return [
            StructuredTool.from_function(
                func=self.create_directory,
                name="create_directory",
                description="Creates a new directory (folder) within the workspace. Useful for organizing files.",
                args_schema=CreateDirectorySchema
            ),
            StructuredTool.from_function(
                func=self.list_directory,
                name="list_directory",
                description="Lists files and directories with their sizes. Useful for exploring content.",
                args_schema=ListDirectorySchema
            )
        ]

class OCRTools(_WorkspaceMixin):
    def __init__(self, root_dir: str, api_key: str):
        self.root_dir = root_dir
        self.api_key = api_key
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://llm.hpc.pcss.pl/v1"
        )
        self.model = "Nanonets-OCR-s"

    def ocr_image(self, file_path: str) -> str:
        """
        Extracts text from an image file (PNG, JPG, JPEG) using OCR.
        Args:
            file_path: The path to the image file.
        """
        try:
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File {file_path} not found."
            
            # Basic mime check
            mime_type, _ = mimetypes.guess_type(full_path)
            if not mime_type or not mime_type.startswith('image'):
                return f"Error: File {file_path} does not appear to be an image ({mime_type})."

            with open(full_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "Transcribe the text in this image verbatim. Output ONLY the text."},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ]
            )
            return response.choices[0].message.content
        except PermissionError:
            raise
        except Exception as e:
            return f"Error performing OCR: {str(e)}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.ocr_image,
                name="ocr_image",
                description="Extracts text from an image file (PNG, JPG) using OCR. Use this to read scanned documents or text in images."
            )
        ]

class PandocTools(_WorkspaceMixin):
    def __init__(self, root_dir: str):
        self.root_dir = root_dir

    def convert_document(self, source_path: str, output_format: str) -> str:
        """
        Converts a document (e.g. HTML) to another format (e.g. DOCX, PDF) using Pandoc.
        Args:
            source_path: Path to the source file (e.g. 'report.html').
            output_format: Target format extension (e.g. 'docx', 'pdf').
        """
        try:
            full_source = self._get_full_path(source_path)
            if not os.path.exists(full_source):
                 return f"Error: Source file {source_path} not found."
            
            # Construct output filename
            base_name = os.path.splitext(source_path)[0]
            target_filename = f"{base_name}.{output_format}"
            full_target = self._get_full_path(target_filename)
            
            # For HTML to PDF, try wkhtmltopdf first (more reliable than LaTeX)
            if output_format.lower() == 'pdf' and source_path.lower().endswith('.html'):
                try:
                    import subprocess
                    result = subprocess.run(
                        ['wkhtmltopdf', '--encoding', 'utf-8', full_source, full_target],
                        capture_output=True, text=True, timeout=30
                    )
                    if result.returncode == 0:
                        return f"Successfully converted {source_path} to {target_filename}."
                except FileNotFoundError:
                    pass  # wkhtmltopdf not installed, fall through to pypandoc
                except Exception:
                    pass
            
            # Fallback to pypandoc
            if pypandoc is None:
                return "Error: pypandoc module is not installed and wkhtmltopdf not available."
            
            output = pypandoc.convert_file(full_source, output_format, outputfile=full_target)
            return f"Successfully converted {source_path} to {target_filename}."
        except Exception as e:
             return f"Error converting document: {str(e)}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.convert_document,
                name="convert_document",
                description="Converts documents between formats (e.g. HTML to DOCX). Args: source_path (str), output_format (str). Example: source_path='file.html', output_format='docx'."
            )
        ]


class VisionTools(_WorkspaceMixin):
    def __init__(self, root_dir: str, api_key: str, model_name: str = None):
        self.root_dir = root_dir
        self.api_key = api_key
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://llm.hpc.pcss.pl/v1"
        )
        self.model = model_name or "Qwen3-VL-235B-A22B-Instruct"
        self.vision_available = True

    def analyze_image(self, file_path: str, prompt: str = "Opisz szczegółowo ten obraz.") -> str:
        """
        Analyzes an image file using a Vision LLM (Qwen3-VL).
        Args:
            file_path: The name of the image file (e.g., 'chart.png').
            prompt: Question or instruction about the image (in Polish/English).
        """
        # PCSS has no multimodal models - return helpful error
        if not self.vision_available:
            return "Error: Vision tools are currently disabled."
        
        try:
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File {file_path} not found."
            
            # Basic mime check
            mime_type, _ = mimetypes.guess_type(full_path)
            if not mime_type or not mime_type.startswith('image'):
                return f"Error: File {file_path} does not appear to be an image ({mime_type})."

            with open(full_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            error_msg = str(e)
            if "not a multimodal model" in error_msg.lower() or "multimodal" in error_msg.lower():
                return f"Error: Model '{self.model}' does not support image analysis. Please use a vision-capable model (e.g., gpt-4o-mini, gpt-4-turbo) or check available models at /v1/models."
            return f"Error analyzing image: {error_msg}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.analyze_image,
                name="analyze_image",
                description="Analizuje obraz przy użyciu modelu multimodalnego (Qwen3-VL). Używaj do opisywania scen, rozumienia wykresów lub analizy układu dokumentów. Wymaga: file_path (ścieżka do pliku obrazu) oraz prompt (pytanie/polecenie)."
            )
        ]


class ChartTools(_WorkspaceMixin):
    """
    Tools for generating charts and visualizations using matplotlib.
    """
    def __init__(self, root_dir: str = "."):
        self.root_dir = root_dir
    
    def generate_chart(
        self, 
        chart_type: str, 
        data: str, 
        labels: str, 
        file_path: str,
        title: str = "Chart",
        x_label: str = "",
        y_label: str = "",
        colors: str = ""
    ) -> str:
        """
        Generate a chart and save it as PNG/JPG.
        
        Args:
            chart_type: Type of chart - 'bar', 'line', 'pie', 'scatter', 'horizontal_bar'
            data: Comma-separated values (e.g., "10,25,30,15,20")
            labels: Comma-separated labels (e.g., "Jan,Feb,Mar,Apr,May")
            file_path: Output file path (e.g., "charts/sales.png")
            title: Chart title
            x_label: X-axis label (for bar/line charts)
            y_label: Y-axis label (for bar/line charts)
            colors: Optional comma-separated colors (e.g., "#FF6384,#36A2EB,#FFCE56")
        
        Returns:
            Success message with file path or error message.
        """
        try:
            import matplotlib
            matplotlib.use('Agg')  # Non-interactive backend
            import matplotlib.pyplot as plt
            import numpy as np
            
            # Parse data
            try:
                values = [float(x.strip()) for x in data.split(",")]
            except ValueError:
                return "Error: 'data' must be comma-separated numbers (e.g., '10,25,30')"
            
            # Parse labels
            label_list = [x.strip() for x in labels.split(",")]
            if len(label_list) != len(values):
                return f"Error: Number of labels ({len(label_list)}) must match number of data points ({len(values)})"
            
            # Parse colors
            if colors:
                color_list = [x.strip() for x in colors.split(",")]
            else:
                # Default color palette
                color_list = plt.cm.Set3.colors[:len(values)]
            
            # Create figure
            fig, ax = plt.subplots(figsize=(10, 6))
            
            chart_type_lower = chart_type.lower().strip()
            
            if chart_type_lower == "bar":
                bars = ax.bar(label_list, values, color=color_list[:len(values)])
                # Add value labels on bars
                for bar, val in zip(bars, values):
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, 
                           f'{val:.1f}', ha='center', va='bottom', fontsize=9)
                           
            elif chart_type_lower == "horizontal_bar":
                bars = ax.barh(label_list, values, color=color_list[:len(values)])
                for bar, val in zip(bars, values):
                    ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                           f'{val:.1f}', ha='left', va='center', fontsize=9)
                           
            elif chart_type_lower == "line":
                ax.plot(label_list, values, marker='o', linewidth=2, markersize=8, 
                       color=color_list[0] if colors else '#36A2EB')
                ax.fill_between(label_list, values, alpha=0.3)
                
            elif chart_type_lower == "pie":
                ax.pie(values, labels=label_list, autopct='%1.1f%%', 
                      colors=color_list[:len(values)], startangle=90)
                ax.axis('equal')
                
            elif chart_type_lower == "scatter":
                # For scatter, use index as X if no separate X data
                x_vals = range(len(values))
                ax.scatter(x_vals, values, c=color_list[:len(values)], s=100)
                ax.set_xticks(list(x_vals))
                ax.set_xticklabels(label_list)
                
            else:
                return f"Error: Unknown chart_type '{chart_type}'. Use: bar, line, pie, scatter, horizontal_bar"
            
            # Set labels and title
            ax.set_title(title, fontsize=14, fontweight='bold', pad=15)
            if x_label and chart_type_lower != "pie":
                ax.set_xlabel(x_label, fontsize=11)
            if y_label and chart_type_lower != "pie":
                ax.set_ylabel(y_label, fontsize=11)
            
            # Style
            if chart_type_lower != "pie":
                ax.spines['top'].set_visible(False)
                ax.spines['right'].set_visible(False)
                ax.grid(axis='y', alpha=0.3)
            
            plt.tight_layout()
            
            # Resolve and validate path through sandbox
            full_path = _safe_path(self.root_dir, file_path)
            os.makedirs(os.path.dirname(full_path) if os.path.dirname(full_path) else ".", exist_ok=True)
            
            # Determine format from extension
            ext = os.path.splitext(file_path)[1].lower()
            if ext not in ['.png', '.jpg', '.jpeg', '.svg', '.pdf']:
                full_path += '.png'
            
            plt.savefig(full_path, dpi=150, bbox_inches='tight', 
                       facecolor='white', edgecolor='none')
            plt.close(fig)
            
            return f"Chart saved successfully: {file_path}"
            
        except ImportError:
            return "Error: matplotlib is required. Install with: pip install matplotlib"
        except PermissionError:
            raise
        except Exception as e:
            return f"Error generating chart: {str(e)}"
    
    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.generate_chart,
                name="generate_chart",
                description="""Generate a chart/visualization and save as PNG/JPG.
Args:
- chart_type: 'bar', 'line', 'pie', 'scatter', or 'horizontal_bar'
- data: Comma-separated values (e.g., "10,25,30,15")
- labels: Comma-separated labels (e.g., "Q1,Q2,Q3,Q4")
- file_path: Output path (e.g., "chart.png")
- title: Chart title
- x_label, y_label: Optional axis labels
- colors: Optional comma-separated hex colors"""
            )
        ]


class WebSearchTools:
    """
    Optimized web search tools for deep research.
    """
    def __init__(self, api_key: str = None, model_name: str = "gpt-4o", base_url: str = "https://llm.hpc.pcss.pl/v1"):
        self.api_key = api_key
        self.model_name = model_name
        self.base_url = base_url
        self.client = None
        if self.api_key:
            try:
                self.client = OpenAI(api_key=self.api_key, base_url=self.base_url)
            except Exception:
                pass

    def search_web(self, query: str, max_results: int = 10) -> str:
        """
        Performs a general web search using DuckDuckGo.
        Best for: definitions, how-to guides, reference information.
        Args:
            query: The search query string.
            max_results: Maximum number of results (default 10).
        """
        try:
            results = []
            with DDGS() as ddgs:
                ddgs_gen = ddgs.text(query, region="pl-pl", max_results=max_results)
                if ddgs_gen:
                    for i, r in enumerate(ddgs_gen, 1):
                        results.append(
                            f"[{i}] {r.get('title', 'No title')}\n"
                            f"    URL: {r.get('href', '')}\n"
                            f"    {r.get('body', '')[:200]}"
                        )
            
            if not results:
                return "No results found. Try rephrasing your query."
            
            return "SEARCH RESULTS:\n\n" + "\n\n".join(results) + "\n\n[TIP: Use visit_page on URLs to read full content]"
        except Exception as e:
            return f"Search error: {str(e)}"

    def search_news(self, query: str, max_results: int = 8) -> str:
        """
        Searches for recent NEWS articles using DuckDuckGo News.
        Best for: current events, breaking news, recent developments.
        Args:
            query: The news search query.
            max_results: Maximum number of news articles (default 8).
        """
        try:
            results = []
            with DDGS() as ddgs:
                news_gen = ddgs.news(query, region="pl-pl", max_results=max_results)
                if news_gen:
                    for i, r in enumerate(news_gen, 1):
                        date = r.get('date', 'Unknown date')
                        source = r.get('source', 'Unknown source')
                        results.append(
                            f"[{i}] {r.get('title', 'No title')}\n"
                            f"    Source: {source} | Date: {date}\n"
                            f"    URL: {r.get('url', '')}\n"
                            f"    {r.get('body', '')[:200]}"
                        )
            
            if not results:
                return "No news found. Try broader search terms."
            
            return "NEWS RESULTS:\n\n" + "\n\n".join(results) + "\n\n[TIP: Use visit_page on URLs to read full articles]"
        except Exception as e:
            return f"News search error: {str(e)}"

    def _summarize_content(self, text: str, max_chars: int = 2000) -> str:
        """
        Uses the LLM to summarize long content into key findings.
        """
        if not self.client or not text:
            return text[:max_chars] + "..." if len(text) > max_chars else text
            
        try:
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "system", "content": "Jesteś ekspertem od analizy danych. Podsumuj poniższy tekst w punktach, wyciągając najważniejsze fakty i wnioski. Odpowiadaj po polsku."},
                    {"role": "user", "content": f"Tekst do podsumowania (maksymalnie {max_chars} znaków):\n\n{text[:6000]}"}
                ],
                temperature=0.3,
                max_tokens=800
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"(Błąd sumaryzacji: {e}) " + text[:max_chars]

    def deep_research(self, topic: str, max_sources: int = 3) -> str:
        """
        Performs comprehensive research on a topic by visiting multiple sources and summarizing findings.
        Args:
            topic: The research topic or question.
            max_sources: Number of top sources to analyze deeply (default 3).
        """
        try:
            # 1. Search for sources
            search_query = topic
            with DDGS() as ddgs:
                results = list(ddgs.text(search_query, region="pl-pl", max_results=max_sources + 2))
            
            if not results:
                return "Nie znaleziono żadnych wyników dla podanego tematu."
            
            research_report = [f"# Raport Deep Research: {topic}\n"]
            
            # 2. Visit and summarize selected sources
            processed_count = 0
            for i, res in enumerate(results):
                if processed_count >= max_sources:
                    break
                    
                url = res.get('href')
                title = res.get('title', 'Brak tytułu')
                
                if not url: continue
                
                content = self.visit_page(url)
                if "Error" in content or "Timeout" in content:
                    continue
                
                summary = self._summarize_content(content)
                research_report.append(f"## [{i+1}] {title}")
                research_report.append(f"**Źródło:** {url}")
                research_report.append(f"**Kluczowe ustalenia:**\n{summary}\n")
                processed_count += 1
            
            if processed_count == 0:
                return "Udało się znaleźć linki, ale nie udało się pobrać treści z żadnego ze źródeł."
                
            return "\n".join(research_report)
            
        except Exception as e:
            return f"Błąd podczas głębokiego researchu: {str(e)}"

    def visit_page(self, url: str) -> str:
        """
        Visits a URL and extracts the main article content.
        Uses readability algorithm to extract relevant text, ignoring navigation/ads.
        Args:
            url: The URL to visit.
        """
        try:
            import requests
            import time
            from readability import Document
            
            headers = {
                "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                "Accept-Language": "pl-PL,pl;q=0.9,en;q=0.8",
            }
            
            # Add retry logic for reliability
            max_retries = 2
            for attempt in range(max_retries):
                try:
                    response = requests.get(url, headers=headers, timeout=20)
                    response.raise_for_status()
                    break
                except (requests.exceptions.RequestException, requests.exceptions.Timeout) as e:
                    if attempt == max_retries - 1:
                        raise e
                    time.sleep(1) # Small delay before retry
            
            # Encoding detection
            content_type = response.headers.get('Content-Type', '')
            if 'charset=' in content_type.lower():
                declared_encoding = content_type.split('charset=')[-1].split(';')[0].strip()
                response.encoding = declared_encoding
            else:
                response.encoding = 'utf-8'
            
            # Use readability to extract main content
            doc = Document(response.text)
            title = doc.title()
            
            from bs4 import BeautifulSoup
            content_html = doc.summary()
            soup = BeautifulSoup(content_html, 'html.parser')
            
            # Extract text with improved structure
            text_parts = []
            for elem in soup.find_all(['p', 'h1', 'h2', 'h3', 'li', 'td', 'th']):
                text = elem.get_text(strip=True)
                if text and len(text) > 15:
                    text_parts.append(text)
            
            content = "\n\n".join(text_parts)
            
            # High-quality fallback
            if not content or len(content) < 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                # Prioritize semantic tags
                main_content = soup.find(['article', 'main']) or soup.find('div', class_=re.compile(r'content|article|post', re.I))
                if main_content:
                    content = main_content.get_text(separator='\n\n', strip=True)
                else:
                    for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside', 'form']):
                        tag.decompose()
                    content = soup.get_text(separator='\n\n', strip=True)
            
            # Higher limit for deep research (15k chars)
            max_limit = 15000
            if len(content) > max_limit:
                 content = content[:max_limit] + "\n\n[... Treść ucięta ze względu na limit ...]"
            
            return f"=== {title} ===\nSource: {url}\n\n{content}"
            
        except ImportError as ie:
            return f"Brak biblioteki: {ie}. Zainstaluj: pip install readability-lxml requests beautifulsoup4"
        except Exception as e:
            return f"Błąd podczas odwiedzania strony {url}: {str(e)}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.search_web,
                name="search_web",
                description="Wyszukuje informacje w internecie. Zwraca linki i krótkie fragmenty. Użyj visit_page aby przeczytać pełną treść."
            ),
            StructuredTool.from_function(
                func=self.search_news,
                name="search_news",
                description="Wyszukuje najnowsze WIADOMOŚCI. Zwraca artykuły z datami. Użyj visit_page aby przeczytać całość."
            ),
            StructuredTool.from_function(
                func=self.visit_page,
                name="visit_page",
                description="Odwiedza podany adres URL i wyciąga czysty tekst z artykułu. Używaj po search_web."
            ),
            StructuredTool.from_function(
                func=self.deep_research,
                name="deep_research",
                description="Automatyczny proces głębokiego researchu. Odwiedza wiele źródeł, podsumowuje je przez AI i tworzy raport zbiorczy. Najlepsze dla złożonych pytań."
            )
        ]


class PythonREPL:
    """
    A simple Python REPL for executing code safely.
    """
    def __init__(self, root_dir: str = "."):
        self.root_dir = root_dir

    def run_python(self, code: str) -> str:
        """
        Executes Python code and returns the standard output.
        Args:
            code: The Python code to execute.
        """
        # --- Hardened REPL Sandbox ---
        # Block dangerous builtins that can escape the workspace.
        blocked_builtins = {'open', 'eval', 'compile',
                            'breakpoint', 'input', 'memoryview'}
        safe_builtins = {k: v for k, v in __builtins__.items()
                         if k not in blocked_builtins} if isinstance(__builtins__, dict) else {
            k: getattr(__builtins__, k) for k in dir(__builtins__)
            if k not in blocked_builtins and not k.startswith('_')
        }
        # Provide safe file I/O restricted to workspace
        root = os.path.realpath(self.root_dir)

        def _safe_open(path, mode='r', *args, **kwargs):
            resolved = os.path.realpath(os.path.join(root, str(path).lstrip('/\\')))
            if not resolved.startswith(root):
                raise PermissionError(f"Access denied: '{path}' is outside the workspace.")
            return open(resolved, mode, *args, **kwargs)

        # A minimal os-like namespace — no os.system, no subprocess, no chdir
        import math, json, datetime as dt, collections, itertools
        try:
            import numpy as np_mod
        except ImportError:
            np_mod = None
        try:
            import pandas as pd_mod
        except ImportError:
            pd_mod = None
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt_mod
            
            # Monkey-patch savefig to ensure files are saved in root_dir
            _orig_savefig = plt_mod.savefig
            def _safe_savefig(*args, **kwargs):
                if args and isinstance(args[0], (str, bytes, os.PathLike)):
                    path = str(args[0])
                    if not os.path.isabs(path):
                        # Prepend root_dir if it's a relative path
                        args = (os.path.join(root, path),) + args[1:]
                return _orig_savefig(*args, **kwargs)
            plt_mod.savefig = _safe_savefig

        except ImportError:
            plt_mod = None

        sandbox_globals = {
            '__builtins__': safe_builtins,
            'open': _safe_open,
            'os': os,
            'math': math,
            'json': json,
            'datetime': dt,
            'collections': collections,
            'itertools': itertools,
            'np': np_mod,
            'pd': pd_mod,
            'plt': plt_mod,
            'matplotlib': plt_mod,
        }

        stdout = io.StringIO()
        stderr = io.StringIO()
        with contextlib.redirect_stdout(stdout), contextlib.redirect_stderr(stderr):
            try:
                exec(code, sandbox_globals)  # noqa: S102
            except Exception as e:
                return f"Error executing Python: {str(e)}"

        output = stdout.getvalue()
        err = stderr.getvalue()
        if err:
            output += f"\n[stderr]:\n{err}"
        return output if output else "Code executed successfully (no output)."

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.run_python,
                name="run_python",
                description="Executes Python code in a local environment. Useful for data processing, math, or complex logic. Returns standard output."
            )
        ]
class ViewFileSchema(BaseModel):
    file_path: str = Field(description="The path to the file to view.")
    start_line: Optional[int] = Field(default=None, description="Optional starting line number (1-indexed).")
    end_line: Optional[int] = Field(default=None, description="Optional ending line number (inclusive).")

class ViewFileTool(_WorkspaceMixin):
    def __init__(self, root_dir: str):
        self.root_dir = root_dir

    def view_file(self, file_path: str, start_line: Optional[int] = None, end_line: Optional[int] = None) -> str:
        """
        Views the contents of a file, prepending 1-indexed line numbers.
        """
        try:
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File {file_path} not found."

            with open(full_path, "r", encoding="utf-8") as f:
                lines = f.readlines()

            if not lines:
                return f"File {file_path} is empty."

            # Setup bounds
            total_lines = len(lines)
            start_idx = max(0, start_line - 1) if start_line else 0
            
            # Truncation logic: If end_line is not provided and file is > 150 lines, truncate.
            if end_line:
                end_idx = min(total_lines, end_line)
                truncated = False
            else:
                if total_lines - start_idx > 150:
                    end_idx = start_idx + 150
                    truncated = True
                else:
                    end_idx = total_lines
                    truncated = False

            # Guard against invalid bounds
            if start_idx >= total_lines or start_idx > end_idx:
                return f"Error: Invalid line range {start_line}-{end_line} for file with {total_lines} lines."

            # Construct numbered output
            output_lines = [f"{i + 1}: {lines[i]}" for i in range(start_idx, end_idx)]
            
            result = f"File: {file_path} (Lines {start_idx + 1}-{end_idx} of {total_lines})\n"
            result += "-" * 40 + "\n"
            result += "".join(output_lines)
            if not result.endswith("\n"):
                 result += "\n"
            result += "-" * 40
            
            if truncated:
                result += f"\n[WARNING: File truncated at 150 lines to prevent context overflow."
                result += f"\n To see the rest of the file, call view_file again with start_line={end_idx + 1} and end_line={min(total_lines, end_idx + 150)}]"
                
            return result
            
        except UnicodeDecodeError:
            return f"Error: {file_path} appears to be a binary file."
        except Exception as e:
            return f"Error viewing file: {str(e)}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.view_file,
                name="view_file",
                description="View file contents with line numbers. CRITICAL: Use this BEFORE edit_file to get the exact line numbers to change.",
                args_schema=ViewFileSchema
            )
        ]


class ReplaceFileContentSchema(BaseModel):
    file_path: str = Field(description="The target file to modify.")
    start_line: int = Field(description="The starting line number of the chunk to replace (1-indexed).")
    end_line: int = Field(description="The ending line number of the chunk (1-indexed, inclusive).")
    replacement_content: str = Field(description="The exact text to replace the specified line range with.")

class ReplaceFileContentTool(_WorkspaceMixin):
    def __init__(self, root_dir: str):
        self.root_dir = root_dir

    def replace_file_content(self, file_path: str, start_line: int, end_line: int, replacement_content: str) -> str:
        """
        Replaces a specific range of lines in a file.
        """
        try:
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File {file_path} not found."

            with open(full_path, "r", encoding="utf-8") as f:
                lines = f.readlines()

            total_lines = len(lines)
            
            # Validation
            if start_line < 1 or end_line < start_line or start_line > total_lines:
                return f"Error: Invalid line range [{start_line}, {end_line}] for file with {total_lines} lines."
                
            # Convert to 0-indexed indices for lists
            start_idx = start_line - 1
            # end_line is inclusive, so the slice goes up to end_line
            end_idx = min(end_line, total_lines)
            
            # Prepare replacement block (ensure it ends with a newline if the file does)
            if replacement_content and not replacement_content.endswith("\n"):
                 replacement_content += "\n"

            # Construct new content
            new_lines = lines[:start_idx] + [replacement_content] + lines[end_idx:]
            
            with open(full_path, "w", encoding="utf-8") as f:
                f.writelines(new_lines)
                
            return f"Successfully replaced lines {start_line}-{end_line} in {file_path}."
        except PermissionError:
            return f"Error: Permission denied writing to {file_path}."
        except Exception as e:
            return f"Error editing file: {str(e)}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.replace_file_content,
                name="replace_file_content",
                description="Modifies a file by replacing a block of lines (start_line to end_line) with new content. REPLACES EXACT STRING MATCHING. Always use view_file first to find the target lines.",
                args_schema=ReplaceFileContentSchema
            )
        ]

class SearchTools(_WorkspaceMixin):
    def __init__(self, root_dir: str):
        self.root_dir = root_dir

    def search_files(self, query: str, pattern: str = "*", recursive: bool = True) -> str:
        """
        Searches for a string (query) inside files matching the pattern.
        Args:
            query: The text to search for.
            pattern: Glob pattern for files (e.g., '*.py').
            recursive: Whether to search subdirectories.
        """
        try:
            results = []
            import fnmatch
            
            search_path = self.root_dir
            for root, dirs, files in os.walk(search_path):
                if not recursive and root != search_path:
                    continue
                
                for filename in fnmatch.filter(files, pattern):
                    if filename.startswith('.'): continue
                    
                    full_path = os.path.join(root, filename)
                    rel_path = os.path.relpath(full_path, self.root_dir)
                    
                    try:
                        with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                            for i, line in enumerate(f, 1):
                                if query.lower() in line.lower():
                                    results.append(f"{rel_path}:{i}: {line.strip()}")
                                    if len(results) > 50:
                                        return "Too many results (truncated):\n" + "\n".join(results[:50])
                    except:
                        continue
            
            if not results:
                return f"No matches found for '{query}'."
            
            return "\n".join(results)
        except PermissionError:
            raise
        except Exception as e:
            return f"Error searching files: {str(e)}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.search_files,
                name="search_files",
                description="Searches for a specific string within files in the workspace. Returns file names and line numbers."
            )
        ]


class TerminalSchema(BaseModel):
    command: str = Field(description="The shell command to execute in the workspace.")

class TerminalTool(_WorkspaceMixin):
    def __init__(self, root_dir: str):
        self.root_dir = os.path.abspath(root_dir)

    def run_terminal(self, command: str) -> str:
        """
        Executes a shell command within the workspace directory.
        Args:
            command: The command to run (e.g., 'python app.py' or 'ls -la').
        """
        try:
            import platform
            is_unix = platform.system() != "Windows"
            
            start_time = time.time()
            process = subprocess.Popen(
                command,
                shell=True,
                cwd=self.root_dir,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                env=os.environ.copy(),
                start_new_session=is_unix  # Put in new process group on Unix
            )
            
            try:
                stdout, stderr = process.communicate(timeout=15) # Shorter timeout for interactive agents
            except subprocess.TimeoutExpired:
                # Command took too long. We need to kill it and all its children.
                if is_unix:
                    import signal
                    try:
                        os.killpg(os.getpgid(process.pid), signal.SIGKILL)
                    except Exception:
                        process.kill()
                else:
                    process.kill()
                    
                # Read whatever was produced before we killed it
                try:
                    stdout, stderr = process.communicate(timeout=2)
                except subprocess.TimeoutExpired:
                    stdout, stderr = "Output extraction hung", "Output extraction hung"
                    
                return f"Command execution stopped after 15s timeout (background server/process detected).\n[Partial Stdout]:\n{stdout}\n[Partial Stderr]:\n{stderr}"

            output = stdout
            if stderr:
                output += f"\n[stderr]:\n{stderr}"
            
            if not output:
                return f"Command executed successfully (exit code {process.returncode}), no output."
            
            # Truncate extremely long output
            if len(output) > 10000:
                output = output[:5000] + "\n... [Output truncated for length] ...\n" + output[-5000:]
                
            return output
        except Exception as e:
            return f"Error executing terminal command: {str(e)}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.run_terminal,
                name="run_terminal",
                description="Executes a shell command in the local workspace terminal. Use this to run applications, tests, or system utilities. DANGER: Use surgical commands.",
                args_schema=TerminalSchema
            )
        ]

class UpdateContextSchema(BaseModel):
    summary: str = Field(description="A concise summary of current project state, key changes, and next steps.")

class UpdateContextTool(_WorkspaceMixin):
    def __init__(self, root_dir: str):
        self.root_dir = os.path.abspath(root_dir)

    def update_context(self, summary: str) -> str:
        """
        Updates the .agent_context.md file in the workspace to maintain persistence across sessions.
        Includes a simple de-duplication filter to prevent runaway repetitive AI output.
        """
        try:
            # De-duplication filter: split by double newlines or lines
            # and keep only unique paragraphs/lines while preserving order.
            lines = summary.split('\n')
            unique_lines = []
            seen = set()
            for line in lines:
                clean_line = line.strip()
                # We only de-duplicate lines that are likely part of repetitive headers/sections
                # and are not empty. 
                if clean_line and (clean_line.startswith('#') or len(clean_line) > 20):
                    if clean_line not in seen:
                        unique_lines.append(line)
                        seen.add(clean_line)
                    else:
                        # Skip repetitive line
                        pass
                else:
                    unique_lines.append(line)
            
            clean_summary = '\n'.join(unique_lines)
            
            context_path = os.path.join(self.root_dir, ".agent_context.md")
            
            with open(context_path, "w", encoding="utf-8") as f:
                import datetime as dt_internal
                timestamp = dt_internal.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                f.write(f"# Project Context\n\n**Last Updated:** {timestamp}\n\n{clean_summary}\n")
                
            msg = "Successfully updated .agent_context.md."
            if len(clean_summary) < len(summary) * 0.8:
                msg += " (Note: redundant repetitive content was automatically filtered)"
            return msg
        except Exception as e:
            return f"Error updating context: {str(e)}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.update_context,
                name="update_context",
                description="Updates a hidden '.agent_context.md' file with the latest project status. USE THIS to 'save your place' for future sessions.",
                args_schema=UpdateContextSchema
            )
        ]
